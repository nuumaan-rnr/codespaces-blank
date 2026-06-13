"""Design Validation Report as editable DOCX and PDF.

The report content is built once as a list of blocks (headings, paragraphs,
tables, images, verdict) and rendered to either format, so DOCX and PDF
stay identical.  The layout follows the EN 15512 worked-example calc-sheet
style: a title block, numbered sections, parameter tables, dimensioned
model views, and per-check verifications citing the EN clause with a
PASS / FAIL ("=> o.k.") statement.
"""

from __future__ import annotations

import datetime as _dt
import io
import os
from typing import Dict, List, Optional, Tuple

import matplotlib.pyplot as plt

from .checks.en15512 import CheckResult, all_ok, governing
from .report_html import CLAUSES        # shared EN clause + basis text
from .viewer import (plot_front_elevation, plot_model, plot_plan,
                     plot_side_elevation, plot_utilization)

CHECK_ORDER = ["STRESS", "BUCKLING", "BRACE_BUCKLING", "CONNECTOR",
               "BRACE_BOLT", "BASEPLATE", "BASE_RESTRAINT", "ANCHORAGE",
               "SPLICE", "DEFLECTION", "SWAY", "ALPHA_CR", "STABILITY"]


def _png(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


# ------------------------------------------------------------- content model
def build_report_blocks(model, cases, checks, meta=None) -> List[tuple]:
    """Return the report as a list of (kind, *payload) blocks."""
    from . import branding as B
    meta = meta or {}
    b: List[tuple] = []
    verdict = all_ok(checks)
    gov = governing(checks)

    logo = B.logo_bytes()
    if logo:
        b.append(("logo", logo, f"{B.COMPANY}  -  {B.TAGLINE}  ({B.WEBSITE})"))
    b.append(("title", "Design Validation Report",
              "Selective pallet racking - EN 15512 (second-order elastic, "
              "semi-rigid connections)"))
    rows = []
    for k in ("project", "system", "configuration", "client", "location",
              "engineer"):
        if meta.get(k):
            rows.append((k.title(), str(meta[k])))
    rows.append(("Standard", "EN 15512 (non-seismic); EN 1993-1-1, -1-3, -1-8"))
    rows.append(("Analysis", "OpenSees - 3D geometrically nonlinear (P-Delta) "
                             "elastic"))
    rows.append(("Date", f"{_dt.datetime.now():%Y-%m-%d %H:%M}"))
    b.append(("kv", rows))
    b.append(("verdict", "OVERALL RESULT: " + ("PASS" if verdict else "FAIL"),
              verdict))
    if gov:
        b.append(("p", f"Governing check: {gov.check} on {gov.target} "
                       f"({gov.member_set}) in {gov.case} - utilisation "
                       f"{gov.utilization:.3f}."))

    # 1 model summary
    b.append(("h2", "1  Model summary"))
    xs = [n.x for n in model.nodes.values()]
    ys = [n.y for n in model.nodes.values()]
    zs = [n.z for n in model.nodes.values()]
    sets = {}
    for m in model.members.values():
        sets[m.member_set] = sets.get(m.member_set, 0) + 1
    b.append(("kv", [
        ("Nodes / members", f"{len(model.nodes)} / {len(model.members)}"),
        ("Overall L x D x H [mm]",
         f"{max(xs)-min(xs):.0f} x {max(ys)-min(ys):.0f} x {max(zs)-min(zs):.0f}"),
        ("Member groups",
         ", ".join(f"{k}: {v}" for k, v in sorted(sets.items()))),
    ]))
    b.append(("h3", "1.1  Sections"))
    srows = [["Section", "Role", "fy", "A", "A_eff", "Iy", "Iz", "J"]]
    for s in model.sections.values():
        fy = model.materials[s.material].fy
        srows.append([s.name, s.role, f"{fy:.0f}", f"{s.A:.0f}",
                      f"{s.area_eff:.0f}", f"{s.Iy:.3e}", f"{s.Iz:.3e}",
                      f"{s.J:.3e}"])
    b.append(("table", srows))

    # 2 supports & stiffness
    b.append(("h2", "2  Supports and stiffness modelling"))
    b.append(("p", "The base (floor-connection) rotational stiffness and the "
                   "beam-to-upright connector stiffness are explicitly "
                   "included in the analysis model (EN 15512 8.4)."))
    b.append(("h3", "2.1  Base / floor connections"))

    def fmt(v):
        return "fixed" if v is True else ("free" if v is False
                                          else f"{float(v)/1e6:.1f}")
    rrows = [["Node", "uX,uY,uZ", "k_rx [kNm/rad]", "k_ry [kNm/rad]", "rz"]]
    for s in model.supports:
        rrows.append([str(s.node), f"{fmt(s.ux)},{fmt(s.uy)},{fmt(s.uz)}",
                      fmt(s.rx), fmt(s.ry), fmt(s.rz)])
    b.append(("table", rrows))
    if model.base_plate and model.base_plate.m_rd_n:
        b.append(("note", "Base moment resistance M_Rd(N) from the "
                          "floor-connection tests is included for the "
                          "partial-restraint check (EN 15512 10.5.1)."))
    b.append(("h3", "2.2  Beam-to-upright connectors (per level)"))
    seen = {}
    for m in model.members.values():
        if m.member_set == "pallet beams" and m.hinge_i:
            seen.setdefault((round(model.nodes[m.node_i].z), m.section),
                            m.hinge_i)
    crows = [["Level z [mm]", "Beam", "k_phi [kNm/rad]", "M_Rd [kNm]",
              "looseness [mrad]"]]
    for (z, sec), h in sorted(seen.items()):
        crows.append([str(z), sec, f"{(h.rz or 0)/1e6:.1f}",
                      f"{(h.m_rd_z or 0)/1e6:.2f}",
                      f"{(h.looseness or 0)*1e3:.1f}"])
    b.append(("table", crows))
    b.append(("note", "The connector rotational stiffness k_phi and the beam "
                      "flexural stiffness EI of the selected section are both "
                      "included in the global analysis."))

    # 3 loads
    b.append(("h2", "3  Load cases and combinations"))
    lrows = [["Load case", "Type"]]
    for lc in model.load_cases.values():
        lrows.append([lc.name, lc.case_type])
    b.append(("table", lrows))
    corows = [["Combination", "Kind", "Load factors", "Imperfection"]]
    for c in model.combinations:
        fac = " + ".join(f"{f:g}*{lc}" for lc, f in c.factors.items())
        imp = (", ".join(c.imp_directions or model.imperfection.directions)
               if c.imperfection else "-")
        corows.append([c.name, c.kind, fac, imp])
    b.append(("table", corows))
    try:
        phi = model.imperfection.value()
        b.append(("p", f"Global sway imperfection phi = {phi:.5f} rad "
                       f"(1/{1/phi:.0f}), method {model.imperfection.method}, "
                       f"applied in "
                       f"{', '.join(model.imperfection.directions)}."))
    except ValueError:
        pass

    # 4 model views
    b.append(("h2", "4  Model views (dimensions in mm)"))
    b.append(("image", _png(plot_front_elevation(model)),
              "Front elevation (down-aisle, X-Z)"))
    b.append(("image", _png(plot_side_elevation(model)),
              "Side elevation (cross-aisle, Y-Z)"))
    b.append(("image", _png(plot_plan(model)), "Plan (top, X-Y)"))
    b.append(("image", _png(plot_model(model)), "3D view"))

    # 5 analysis cases
    b.append(("h2", "5  Analysis cases - convergence"))
    arows = [["Case", "Kind", "Converged", "Sway X", "Sway Y", "alpha_cr"]]
    for c in cases:
        acr = c.alpha_cr_estimate
        arows.append([c.name, c.kind, "yes" if c.converged else "NO",
                      f"{c.max_sway_x:.2f}", f"{c.max_sway_y:.2f}",
                      f"{acr:.1f}" if acr else "-"])
    b.append(("table", arows))

    # 6 design checks
    b.append(("h2", "6  EN 15512 design verifications"))
    b.append(("image", _png(plot_utilization(model, checks)),
              "Governing member utilisation (red > 1 fails)"))
    n = 1
    for kind in CHECK_ORDER:
        crows = [c for c in checks if c.check == kind]
        if not crows:
            continue
        clause, basis = CLAUSES.get(kind, ("", ""))
        crows.sort(key=lambda c: -c.utilization)
        worst = crows[0]
        grp_ok = all(c.ok for c in crows)
        b.append(("h3", f"6.{n}  {kind} - {clause}"))
        n += 1
        b.append(("note", basis))
        b.append(("result", f"{'PASS => o.k.' if grp_ok else 'FAIL => NOT o.k.'}"
                            f"  (worst utilisation {worst.utilization:.3f} on "
                            f"{worst.target}, {worst.case})", grp_ok))
        trows = [["Target", "Set", "Case", "Util.", "Status", "Detail"]]
        for c in crows[:25]:
            trows.append([c.target, c.member_set, c.case,
                          f"{c.utilization:.3f}", c.status, c.detail])
        if len(crows) > 25:
            trows.append([f"... {len(crows)-25} more rows", "", "", "", "", ""])
        b.append(("table", trows))

    b.append(("note", f"Generated by {B.COMPANY} {B.PRODUCT} "
                      f"({B.WEBSITE}). This report is an engineering aid. "
                      "All partial factors, imperfection parameters and "
                      "section / connector / floor-connection test values "
                      "must be verified by a qualified engineer against the "
                      "applicable edition of EN 15512 and the national "
                      "annex. Scope: non-seismic."))
    return b


# --------------------------------------------------------------- DOCX render
def _add_page_field(paragraph):
    """Insert a Word PAGE field (live page number) into a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addprevious(fld)


def _build_header_footer(doc):
    """Small top-left logo header (pages 2+) and a company / page no /
    copyright footer on every page."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
    from . import branding as B

    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    usable = sec.page_width - sec.left_margin - sec.right_margin

    # header logo on continuation pages (first page has the big body logo)
    logo = B.logo_bytes()
    if logo:
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            hp.add_run().add_picture(io.BytesIO(logo), width=Inches(1.4))
        except Exception:
            pass

    from docx.shared import Emu

    def footer_para(par):
        par.paragraph_format.tab_stops.add_tab_stop(
            Emu(int(usable / 2)), WD_TAB_ALIGNMENT.CENTER)
        par.paragraph_format.tab_stops.add_tab_stop(
            Emu(int(usable)), WD_TAB_ALIGNMENT.RIGHT)
        r = par.add_run(f"{B.COMPANY}\tPage ")
        r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        _add_page_field(par)
        r2 = par.add_run(f"\t© {B.COMPANY} · {B.WEBSITE}")
        r2.font.size = Pt(7.5)
        r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    footer_para(sec.footer.paragraphs[0])
    footer_para(sec.first_page_footer.paragraphs[0])


def render_docx(blocks, path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    TEAL = RGBColor(0x0C, 0x84, 0x90)
    GREY = RGBColor(0x54, 0x54, 0x54)
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9)
    _build_header_footer(doc)

    def shade(par, ok):
        par.runs[0].font.bold = True
        par.runs[0].font.color.rgb = (RGBColor(0x0C, 0x70, 0x80) if ok
                                      else RGBColor(0xB7, 0x1C, 0x1C))

    for blk in blocks:
        kind = blk[0]
        if kind == "logo":
            p = doc.add_paragraph()
            p.add_run().add_picture(io.BytesIO(blk[1]), width=Inches(3.0))
            cap = doc.add_paragraph(blk[2])
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(8)
            cap.runs[0].font.color.rgb = TEAL
        elif kind == "title":
            h = doc.add_heading(blk[1], level=0)
            for r in h.runs:
                r.font.color.rgb = GREY
            doc.add_paragraph(blk[2]).italic = True
        elif kind == "h2":
            h = doc.add_heading(blk[1], level=1)
            for r in h.runs:
                r.font.color.rgb = TEAL
        elif kind == "h3":
            h = doc.add_heading(blk[1], level=2)
            for r in h.runs:
                r.font.color.rgb = GREY
        elif kind == "p":
            doc.add_paragraph(blk[1])
        elif kind == "note":
            p = doc.add_paragraph()
            r = p.add_run(blk[1]); r.italic = True; r.font.size = Pt(8)
        elif kind == "verdict":
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(blk[1]); r.font.size = Pt(15); shade(p, blk[2])
        elif kind == "result":
            p = doc.add_paragraph(); p.add_run(blk[1]); shade(p, blk[2])
        elif kind == "kv":
            t = doc.add_table(rows=0, cols=2); t.style = "Light Grid Accent 1"
            for k, v in blk[1]:
                cells = t.add_row().cells
                cells[0].text, cells[1].text = str(k), str(v)
                cells[0].paragraphs[0].runs[0].font.bold = True
        elif kind == "table":
            data = blk[1]
            t = doc.add_table(rows=1, cols=len(data[0]))
            t.style = "Light Grid Accent 1"
            for j, h in enumerate(data[0]):
                c = t.rows[0].cells[j]; c.text = str(h)
                c.paragraphs[0].runs[0].font.bold = True
            for row in data[1:]:
                cells = t.add_row().cells
                for j, val in enumerate(row):
                    cells[j].text = str(val)
                    for par in cells[j].paragraphs:
                        for r in par.runs:
                            r.font.size = Pt(7)
        elif kind == "image":
            doc.add_picture(io.BytesIO(blk[1]), width=Inches(5.8))
            cap = doc.add_paragraph(blk[2])
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True; cap.runs[0].font.size = Pt(8)
    doc.save(path)
    return path


# ---------------------------------------------------------------- PDF render
def render_pdf(blocks, path):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (Image, Paragraph, SimpleDocTemplate,
                                    Spacer, Table, TableStyle)

    # register a Unicode font (matplotlib ships DejaVuSans)
    import matplotlib
    ttf = os.path.join(os.path.dirname(matplotlib.__file__),
                       "mpl-data", "fonts", "ttf")
    pdfmetrics.registerFont(TTFont("DejaVu", os.path.join(ttf, "DejaVuSans.ttf")))
    pdfmetrics.registerFont(TTFont("DejaVu-Bold",
                                   os.path.join(ttf, "DejaVuSans-Bold.ttf")))
    ss = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=ss["Normal"], fontName="DejaVu",
                          fontSize=8, leading=12, spaceAfter=2,
                          wordWrap="CJK")     # break very long unbroken text
    note = ParagraphStyle("note", parent=body, fontSize=7, leading=10,
                          textColor=colors.grey)
    h1 = ParagraphStyle("h1", parent=body, fontName="DejaVu-Bold",
                        fontSize=18, leading=22, spaceBefore=4, spaceAfter=6,
                        textColor=colors.HexColor("#545454"))
    h2 = ParagraphStyle("h2", parent=body, fontName="DejaVu-Bold",
                        fontSize=12.5, leading=16, spaceBefore=14,
                        spaceAfter=6, textColor=colors.HexColor("#0C8490"))
    h3 = ParagraphStyle("h3", parent=body, fontName="DejaVu-Bold",
                        fontSize=10, leading=14, spaceBefore=8, spaceAfter=3,
                        textColor=colors.HexColor("#545454"))
    cell = ParagraphStyle("cell", parent=body, fontSize=6.8, leading=8.5,
                          spaceAfter=0, wordWrap="CJK")
    cellh = ParagraphStyle("cellh", parent=cell, fontName="DejaVu-Bold")

    story = []
    LM = RM = 15 * mm
    avail = A4[0] - LM - RM

    def add_table(data, widths=None):
        rows = [[Paragraph(str(c), cellh if i == 0 else cell)
                 for c in row] for i, row in enumerate(data)]
        t = Table(rows, colWidths=widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF3F4")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 2.5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, colors.HexColor("#f7f9fb")])]))
        story.append(t)
        story.append(Spacer(1, 6))

    def col_widths(header):
        """Weighted widths so a wide 'Detail' column wraps instead of
        crushing the narrow columns."""
        weights = []
        for h in header:
            hl = str(h).lower()
            weights.append(3.2 if "detail" in hl
                           else 1.6 if hl in ("case", "load factors")
                           else 1.0)
        tot = sum(weights)
        return [avail * w / tot for w in weights]

    brand = ParagraphStyle("brand", parent=note,
                           textColor=colors.HexColor("#0C8490"))
    for blk in blocks:
        kind = blk[0]
        if kind == "logo":
            from PIL import Image as PILImage
            bio = io.BytesIO(blk[1])
            iw, ih = PILImage.open(bio).size
            bio.seek(0)
            w = 70 * mm
            img = Image(bio, width=w, height=w * ih / iw)
            img.hAlign = "LEFT"
            story.append(img)
            story.append(Paragraph("<i>%s</i>" % blk[2], brand))
            story.append(Spacer(1, 10))
        elif kind == "title":
            story.append(Paragraph(blk[1], h1))
            story.append(Paragraph("<i>%s</i>" % blk[2], body))
            story.append(Spacer(1, 8))
        elif kind == "h2":
            story.append(Paragraph(blk[1], h2))
        elif kind == "h3":
            story.append(Paragraph(blk[1], h3))
        elif kind == "p":
            story.append(Paragraph(blk[1], body)); story.append(Spacer(1, 4))
        elif kind == "note":
            story.append(Paragraph("<i>%s</i>" % blk[1], note))
            story.append(Spacer(1, 4))
        elif kind in ("verdict", "result"):
            col = "#0C7080" if blk[2] else "#b71c1c"
            sz = 14 if kind == "verdict" else 9
            story.append(Spacer(1, 2))
            story.append(Paragraph(
                f'<para align="{"center" if kind=="verdict" else "left"}">'
                f'<font color="{col}" size="{sz}"><b>{blk[1]}</b></font></para>',
                body))
            story.append(Spacer(1, 6))
        elif kind == "kv":
            add_table([[k, v] for k, v in blk[1]],
                      widths=[45 * mm, avail - 45 * mm])
        elif kind == "table":
            add_table(blk[1], widths=col_widths(blk[1][0]))
        elif kind == "image":
            bio = io.BytesIO(blk[1])
            from PIL import Image as PILImage
            iw, ih = PILImage.open(bio).size
            bio.seek(0)
            aspect = ih / iw
            w = min(avail, 150 * mm)
            h = w * aspect
            max_h = 190 * mm                 # fit within the page frame
            if h > max_h:
                h, w = max_h, max_h / aspect
            img = Image(bio, width=w, height=h)
            img.hAlign = "CENTER"
            story.append(img)
            story.append(Paragraph("<i>%s</i>" % blk[2], note))
            story.append(Spacer(1, 8))

    from . import branding as B
    logo = B.logo_bytes()
    logo_reader = None
    if logo:
        from reportlab.lib.utils import ImageReader
        logo_reader = ImageReader(io.BytesIO(logo))

    def _footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("DejaVu", 7)
        canvas.setFillColor(colors.HexColor("#888888"))
        y = 9 * mm
        canvas.setStrokeColor(colors.HexColor("#cccccc"))
        canvas.line(LM, y + 4 * mm, A4[0] - RM, y + 4 * mm)
        canvas.drawString(LM, y, B.COMPANY)
        canvas.drawCentredString(A4[0] / 2.0, y, f"Page {doc.page}")
        canvas.drawRightString(A4[0] - RM, y,
                               f"© {B.COMPANY} · {B.WEBSITE}")
        canvas.restoreState()

    def _on_first(canvas, doc):
        _footer(canvas, doc)            # big logo is in the body on page 1

    def _on_later(canvas, doc):
        # small logo, top-left header, on every page after the first
        if logo_reader is not None:
            iw, ih = logo_reader.getSize()
            w = 30 * mm
            h = w * ih / iw
            canvas.drawImage(logo_reader, LM, A4[1] - 11 * mm - h,
                             width=w, height=h, mask="auto")
            canvas.setStrokeColor(colors.HexColor("#0C8490"))
            canvas.setLineWidth(1.0)
            canvas.line(LM, A4[1] - 25 * mm, A4[0] - RM, A4[1] - 25 * mm)
        _footer(canvas, doc)

    # uniform 28 mm top margin leaves room for the later-page header band
    doc = SimpleDocTemplate(path, pagesize=A4, leftMargin=LM, rightMargin=RM,
                            topMargin=28 * mm, bottomMargin=18 * mm,
                            title="Design Validation Report",
                            author=B.COMPANY)
    doc.build(story, onFirstPage=_on_first, onLaterPages=_on_later)
    return path


def write_reports(model, cases, checks, meta, docx_path=None, pdf_path=None):
    blocks = build_report_blocks(model, cases, checks, meta)
    if docx_path:
        render_docx(blocks, docx_path)
    if pdf_path:
        render_pdf(blocks, pdf_path)
    return blocks
